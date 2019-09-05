#!/usr/bin/env python
# coding: utf-8

# In[1]:


from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.common.action_chains import ActionChains

from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from time import sleep
from dateutil import relativedelta
from sortedcontainers import SortedDict, SortedList
import os, sys, re, csv, time, io, filecmp, math
import datetime as dt
import pandas as pd
import configparser as cfgparser


# In[2]:


timestamp = dt.datetime.now().strftime('%H%M')


# In[3]:


cfg = cfgparser.ConfigParser()
cfg.read('C:\\Users\\TFD\\主管差假搜尋\\config.ini')
account = cfg.get('DEFAULT', 'account', fallback = '')
password = cfg.get('DEFAULT', 'password', fallback = '')


# In[4]:


# Frame switch function
def switch_to_iframe(browser):
    browser.switch_to.default_content()
    browser.switch_to.frame("topmost")
    browser.switch_to.frame("iframe")
def switch_to_topmost(browser):
    browser.switch_to.default_content()
    browser.switch_to.frame("topmost")
def switch_to_accountframe(browser):
    browser.switch_to.default_content()
    browser.switch_to.frame("topmost")
    browser.switch_to.frame(browser.find_element_by_css_selector("body > center > table:nth-child(1) > tbody > tr:nth-child(1) > td > table > tbody > tr > td:nth-child(2) > div > iframe"))


# In[5]:


class MyDocument():
    def __init__(self, doc_name):
        self.doc = Document(doc_name)
        
        regex = '(\w+)\t(\S+) (\w+-\w+-\w+)\(\w+\) (\w+:\w+) ~ (\w+-\w+-\w+)\(\w+\) (\w+:\w+)\n\t事由: ([\S ]+) \n\t代理人: (\S+)'
        self.re_comp = re.compile(regex)
        
        # The time limit of the leave period
        td = dt.date.today()
        self.start_lmt = dt.datetime(td.year, td.month, td.day, 8, 30)
        self.mid1_lmt = dt.datetime(td.year, td.month, td.day, 12, 30)
        self.mid2_lmt = dt.datetime(td.year, td.month, td.day, 13, 30)
        self.end_lmt = dt.datetime(td.year, td.month, td.day, 17, 30)
        
        # Paragraph style: 'Normal'
        styles = self.doc.styles
        style = styles['Normal']
        style.font.name = u'標楷體'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
        style.font.size = Pt(28)
        style.font.bold = True
        # New paragraph style: 'New'
        if 'NewParagraph' not in styles:
            newpg = styles.add_style('NewParagraph', WD_STYLE_TYPE.PARAGRAPH)
            newpg.font.name = '標楷體'
            newpg._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
            newpg.font.size = Pt(12)
            newpg.font.bold = True
        # Big font style: 'BigFont'
        if 'BigFont' not in styles:
            bigfont = styles.add_style('BigFont', WD_STYLE_TYPE.CHARACTER)
            bigfont.font.name = '標楷體'
            bigfont._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
            bigfont.font.size = Pt(28)
        # Small font style: 'SmallFont'
        if 'SmallFont' not in styles:
            smallfont = styles.add_style('SmallFont', WD_STYLE_TYPE.CHARACTER)
            smallfont.font.name = '標楷體'
            smallfont._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
            smallfont.font.size = Pt(12)
        
    def setTitleDate(self, date):
        run = self.doc.tables[0].cell(0, 0).paragraphs[0].runs[0]
        run.text = '{}年{}月{}日 主管差假'.format(date.year - 1911, 
                                                 date.month, 
                                                 date.day)
        
    def addParagraph(self):
        newpg = self.doc.tables[0].cell(1, 0).add_paragraph(style = 'NewParagraph')
        newpg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    def writeBigFont(self, string):
        pg = self.doc.tables[0].cell(1, 0).paragraphs[-1]
        pg.add_run(string, style = 'BigFont')
        self.addParagraph()
        
    def writeSmallFont(self, string):
        pg = self.doc.tables[0].cell(1, 0).paragraphs[-1]
        pg.add_run(string, style = 'SmallFont')
        self.addParagraph()
        
    def addLeave(self, raw_string, mode):
        leave, reason, substitute, full_period = self.leaveStringProcess(raw_string)
        self.writeBigFont(leave)
        if mode == 'reason':
            self.writeSmallFont('(' + reason + ')')
        elif mode == 'substitute':
            self.writeSmallFont('(代理人: ' + substitute + ')')
            self.writeSmallFont('(完整請假區間: ' + full_period + ')')
        elif mode == 'all':
            self.writeSmallFont('(' + reason + ')')
            self.writeSmallFont('(完整請假區間: ' + full_period + ')')
    
    def addExtendPeriod(self, raw_string):
        pass
        
    def leaveStringProcess(self, string):
        '''
        string: ex. 游專委家懿\t休假 108-05-03(二) 09:30 ~ 108-05-04(二) 12:30\n\t事由: 國旅卡休假
        return ('游專委家懿9-12時休假', '國旅卡休假')
        '''
        (officername, leavetype, start_date, start_time, 
           end_date, end_time, reason, 
           substitute) = self.re_comp.findall(string)[0]
        # Leave period 
        start_date = self.changeYear(start_date)
        end_date = self.changeYear(end_date)
        start_dt = start_date + ' ' + start_time
        end_dt = end_date + ' ' + end_time
        onduty = dt.datetime.strptime(start_dt, '%Y-%m-%d %H:%M')
        offduty = dt.datetime.strptime(end_dt,  '%Y-%m-%d %H:%M')
        
        # Using '上午', '下午', or hour numbers
        StartFromMorning = False
        StartFromNoon = False
        EndInNoon = False
        EndInAfternoon = False
        
        if (onduty - self.start_lmt).total_seconds() <= 0:
            StartFromMorning = True
        elif onduty.hour in [12, 13]:
            StartFromNoon = True
        if(offduty - self.end_lmt).total_seconds() >= 0:
            EndInAfternoon = True
        elif offduty.hour in [12, 13]:
            EndInNoon = True
            
        if StartFromMorning and EndInAfternoon:
            date_str = ''
        elif StartFromMorning and EndInNoon:
            date_str = '上午'
        elif StartFromNoon and EndInAfternoon:
            date_str = '下午'
        else:
            date_str = ''
            if StartFromMorning:
                date_str += '8'
            else:
                date_str += str(onduty.hour)

            date_str += '-'

            if EndInAfternoon:
                date_str += '17'
            else:
                date_str += str(offduty.hour)

            date_str += '時'
            
        full_period = re.findall('(\d+-\d+-\d+\(\w+\) \d+:\d+ ~ \d+-\d+-\d+\(\w+\) \d+:\d+)', string)[0]
            
        # Change '主任秘書' into '主秘'
        if '主任秘書' in officername:
            officername = officername.replace('主任秘書', '主秘')
        # Change leavetype
        if ('其他假' in leavetype) or ('補休' in leavetype):
            leavetype = '補休'
            
        leave_str = officername + date_str + leavetype
        return leave_str, reason, substitute, full_period
    
    def changeYear(self, date_str):
        '''
        date_str: ex. 108-10-05
        return: ex. 2019-10-05
        '''
        dt_list = date_str.split('-')
        dt_list[0] = str(int(dt_list[0]) + 1911)
        return ('-'.join(dt_list))
        
    def save(self, save_name):
        self.doc.save(save_name)


# In[6]:


# 請假清單
leavedict = {}     #所有假單的dict (title -> name -> leavelist)
# leavelist = []   #已批核/未批核假單


# In[7]:


# Unit list
OfficeList = ["資通作業科", "減災規劃科", "整備應變科", "局長室", "第一副局長室", "災害搶救科", 
              "火災調查科", "綜合企劃科", "緊急救護科", "督察室", "第三副局長室", "主任秘書室", 
              "秘書室", "人事室", "政風室", "會計室", "訓練中心", "救災救護指揮中心", "火災預防科", 
              "第二副局長室"]


# In[8]:


# File output directory
file_dir = r'Z:\\02獎懲待遇股\\湘苹\\主管差假\\'
detail_dir = r'Z:\\02獎懲待遇股\\湘苹\\主管差假\\主管請假清單\\'


# In[9]:


# Month list
monthlist = {'一': '1', '二': '2', '三': '3', '四': '4', '五': '5', 
             '六': '6', '七': '7', '八': '8', '九': '9', '十': '10', 
             '十一': '11', '十二': '12'}


# In[10]:


# Todo List
# windowless
# multiple leaves
# password


# In[11]:


print("正在執行主管差假搜尋，請勿關閉此視窗")


# In[12]:


option = webdriver.ChromeOptions()
# option.add_argument('--headless')
# option.add_argument('--window-size=1280,800"')
if getattr(sys, 'frozen', False) :
    # running in a bundle
    chromedriver_path = os.path.join(sys._MEIPASS, 'chromedriver.exe')
    browser = webdriver.Chrome(chromedriver_path, options=option)
else:
    # executed as a simple script, the driver should be in "PATH"
    browser = webdriver.Chrome(options=option)
#開啟google首頁
browser.get('https://webitr.gov.taipei/WebITR/')


# In[13]:


# Login
element_user = browser.find_element_by_id("userName")
element_user.send_keys(account)
element_pw = browser.find_element_by_id("login_key")
element_pw.send_keys(password)
element_pw.send_keys(Keys.RETURN)


# In[14]:


# 主管名單更新
titlename_dict = {}
titlename_pair = {}
name2office_pair = {}
# Click "差勤管理/制度管理/基本資料維護"
switch_to_topmost(browser)
wait = WebDriverWait(browser, 10)
## 差勤管理
li1 = browser.find_element_by_css_selector("#MenuBar1 > li:nth-child(5) > a")
ActionChains(browser).move_to_element_with_offset(li1, 10, -10).move_to_element_with_offset(li1, 10, 10).perform()
element = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#MenuBar1 > li:nth-child(5) > ul > li:nth-child(1) > a")))
## 制度管理
li2 = browser.find_element_by_css_selector("#MenuBar1 > li:nth-child(5) > ul > li:nth-child(1) > a")
ActionChains(browser).move_to_element_with_offset(li2, 10, -10).move_to_element_with_offset(li2, 10, 10).perform()
element = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#MenuBar1 > li:nth-child(5) > ul > li:nth-child(1) > ul > li:nth-child(3) > a")))
## 基本資料維護
li3 = browser.find_element_by_css_selector("#MenuBar1 > li:nth-child(5) > ul > li:nth-child(1) > ul > li:nth-child(3) > a")
ActionChains(browser).move_to_element_with_offset(li3, 10, -10).move_to_element_with_offset(li3, 10, 10).perform()
li3.click()
## 移出左方列表使浮出列表消失
li4 = browser.find_element_by_css_selector("#MenuBar1 > li:nth-child(1) > a")
ActionChains(browser).move_to_element_with_offset(li4, 10, -10).perform()
element = wait.until(EC.invisibility_of_element_located((By.CSS_SELECTOR, "#MenuBar1 > li:nth-child(5) > ul > li:nth-child(1) > a")))

switch_to_iframe(browser)
unit_procnum = 0
for office_name in OfficeList:
    unit = Select(browser.find_element_by_id("unit"))
    unit.select_by_visible_text(office_name)
    namelist = browser.find_elements_by_name("persons")
    if len(namelist) == 0:
    ## 無人處室
        continue
    ## 局長室全選
    elif office_name == "局長室":
        browser.find_element_by_id("personsAll").click()
    ## 其餘科室
    else:
        namelist[0].click()
    exebtn = browser.find_element_by_name("exeBtn")
    browser.execute_script("arguments[0].scrollIntoView(false);", exebtn)
    exebtn.click()

    rows = browser.find_elements_by_class_name("p4_list_mouse")
    for row in rows:
        name = row.find_elements_by_tag_name("td")[2].text
        title = row.find_elements_by_tag_name("td")[1].text
        office = row.find_elements_by_tag_name("td")[0].text
        if title == "局長室":
            title = "局長"
        elif title == "專門委員":
            title = "專委"
        elif title == "簡任技正":
            title = "簡技"
        title_name = name[0]+title+name[1:]
        titlename_dict[name] = title_name
        titlename_pair[name] = title
        name2office_pair[name] = office
    
    unit_procnum += 1
    print("\r(1/5) 主管名單更新: %3d%%"%(unit_procnum/20*100), end = '')
print("\r(1/5) 主管名單更新:   完成")


# In[15]:


# Construct a name-to-personalID dictionary
print('搜尋主管身分證號')
main_window = browser.window_handles[0]
switch_to_accountframe(browser)
admin_login_btn = browser.find_element_by_id('adminLoginByName')
browser.execute_script("arguments[0].scrollIntoView(false);", admin_login_btn)
admin_login_btn.click()
pop_window = browser.window_handles[1]
browser.switch_to.window(pop_window)
NameIdList = {}
for name in titlename_dict.keys():
    element = wait.until(EC.visibility_of_element_located((By.NAME, "searchName")))
    name_input = browser.find_element_by_name('searchName')
    browser.execute_script("arguments[0].scrollIntoView(false);", name_input)
    name_input.send_keys(name)
    name_input.send_keys(Keys.RETURN)
    try:
        tb = browser.find_elements_by_css_selector('body > div > form > table > tbody')[0]
        trs = tb.find_elements_by_tag_name('tr')
        for i in range(1, len(trs)):
            tr = trs[i]
            if name2office_pair[name] in tr.text:
                NameIdList[name] = tr.find_element_by_tag_name('input')                                     .get_attribute('value')
    except:
        print("搜尋不到{}之身分證字號".format(name))
browser.close()
browser.switch_to.window(main_window)


# In[16]:


# Click "系統維護/表單進度查詢"
switch_to_topmost(browser)
wait = WebDriverWait(browser, 10)
## 系統維護
li1 = browser.find_element_by_css_selector("#MenuBar1 > li:nth-child(9) > a")
ActionChains(browser).move_to_element_with_offset(li1, 10, -10).move_to_element_with_offset(li1, 10, 10).perform()
element = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#MenuBar1 > li:nth-child(9) > ul > li:nth-child(1) > a")))
## 表單進度查詢
li2 = browser.find_element_by_css_selector("#MenuBar1 > li:nth-child(9) > ul > li:nth-child(1) > a")
ActionChains(browser).move_to_element_with_offset(li2, 10, -10).move_to_element_with_offset(li2, 10, 10).perform()
li2.click()
## 移出左方列表使浮出列表消失
li3 = browser.find_element_by_css_selector("#MenuBar1 > li:nth-child(1) > a")
ActionChains(browser).move_to_element_with_offset(li3, 10, -10).perform()
element = wait.until(EC.invisibility_of_element_located((By.CSS_SELECTOR, "#MenuBar1 > li:nth-child(9) > ul > li:nth-child(1) > a")))
switch_to_iframe(browser)


# In[17]:


def chooseDay(chosen_date, browser):
    # Back to that month
    cal_str = browser.find_elements_by_class_name('title')
    cal_yr_mon = re.findall(r'民國\ (\d+)\ 年\ (.*)月', cal_str[-1].text)[0]
    cal_yr = str(int(cal_yr_mon[0])+1911)
    cal_yr_mon = cal_yr + '-' + monthlist[cal_yr_mon[1]]
    mon1 = dt.datetime.strptime(cal_yr_mon, '%Y-%m')
    mon2 = dt.datetime.strptime(
        chosen_date.strftime("%Y-%m"), '%Y-%m')
    pressnum = relativedelta.relativedelta(mon1, mon2).months

    bar = browser.find_elements_by_class_name('headrow')
    if pressnum > 0:
        monthbtn = bar[-1].find_elements_by_class_name('button')[1]
    else:
        monthbtn = bar[-1].find_elements_by_class_name('button')[3]
    browser.execute_script("arguments[0].scrollIntoView(false);", monthbtn)
    for i in range(abs(pressnum)):
        monthbtn.click()
        
    # Click the chosen day of that month
    calendar = browser.find_elements_by_class_name('calendar')
    days = calendar[-1].find_elements_by_class_name('day')
    days_in_month = []
    for day in days:
        if 'name' not in day.get_attribute('class') and            'othermonth' not in day.get_attribute('class'):
            days_in_month.append(day)
    sel_day = days_in_month[chosen_date.day - 1]
    browser.execute_script("arguments[0].scrollIntoView(false);", sel_day)
    sel_day.click()
 


# In[18]:


def searchTomorrow(leave_dt):
    tl = re.findall('(\w+)-(\w+-\w+)\(\w+\) (\w+:\w+) ~ (\w+)-(\w+-\w+)\(\w+\) (\w+:\w+)', leave_dt)[0]
    st_dt = dt.datetime.strptime(tl[1] + tl[2], '%m-%d%H:%M')
    st_dt = st_dt.replace(year = int(tl[0]) + 1911)
    ed_dt = dt.datetime.strptime(tl[4] + tl[5], '%m-%d%H:%M')
    ed_dt = ed_dt.replace(year = int(tl[3]) + 1911)
    search_det1 = dt.datetime(
        dt.date.today().year, 
        dt.date.today().month,
        dt.date.today().day,
        17, 30)
    search_det2 = dt.datetime(
        dt.date.today().year, 
        dt.date.today().month,
        dt.date.today().day,
        8, 30) + dt.timedelta(days = 1)
    if (ed_dt-search_det1).total_seconds() >= 0 and        (ed_dt-search_det2).total_seconds() <= 0:
        return True
    else:
        return False


# In[19]:


def startFromMorning(leave_dt):
    tl = re.findall('(\w+)-(\w+-\w+)\(\w+\) (\w+:\w+) ~ (\w+)-(\w+-\w+)\(\w+\) (\w+:\w+)', leave_dt)[0]
    st_dt = dt.datetime.strptime(tl[1] + tl[2], '%m-%d%H:%M')
    st_dt = st_dt.replace(year = int(tl[0]) + 1911)
    ed_dt = dt.datetime.strptime(tl[4] + tl[5], '%m-%d%H:%M')
    ed_dt = ed_dt.replace(year = int(tl[3]) + 1911)
    morning_det = dt.datetime(
        dt.date.today().year, 
        dt.date.today().month,
        dt.date.today().day,
        8, 30) + dt.timedelta(days = 1)
    if (st_dt-morning_det).total_seconds() == 0:
        return True
    else:
        return False


# In[20]:


# 待/已批核假查詢 new
leavedict = {}
nextdaylist = []
switch_to_iframe(browser)
print("\r(2/5) 已批核/待批核假單查詢:   ", end = '')
for name in NameIdList.keys():
    sel_day = dt.date.today()
    expectedstartdt = None
    ## 選取部門
    depart = Select(browser.find_element_by_id("depart"))
    depart.select_by_visible_text(name2office_pair[name])
    # Deselect all members
    try:
        selAllPerson = browser.find_element_by_id('selAll')
        selAllPerson.click()
        selAllPerson.click()
    except:
        pass
    # Select member
    mem_name = NameIdList[name]
    name_xpath = "//input[@value='" + mem_name + "']"
    mem_box = browser.find_element_by_xpath(name_xpath)
    browser.execute_script("arguments[0].scrollIntoView(false);", mem_box)
    mem_box.click()
    ## 選取請假開始日期
    datestrbtn = browser.find_element_by_id("begintime_dc_bt")
    browser.execute_script("arguments[0].scrollIntoView(false);", datestrbtn)
    datestrbtn.click()
    chooseDay(sel_day, browser)
    ##選取請假結束日期
    dateendbtn = browser.find_element_by_id("endtime_dc_bt")
    browser.execute_script("arguments[0].scrollIntoView(false);", dateendbtn)
    dateendbtn.click()
    chooseDay(sel_day, browser)
    for sstate in ["申請已批核完成", "申請未批核完成"]:
        ## 選取表單狀態
        sheetstate = Select(browser.find_element_by_id("condition"))
        sheetstate.select_by_visible_text(sstate)
        ## 點選查詢
        querybtn = browser.find_element_by_name("selecta")
        browser.execute_script("arguments[0].scrollIntoView(false);", querybtn)
        querybtn.click()
        ## 從搜尋結果中擷取長官差假
        table = browser.find_element_by_css_selector("#div2 > table")
        rows = table.find_elements_by_class_name("stripeMe")
        for row in rows:
            leavetype = row.find_elements_by_tag_name("td")[2].text
            # Find the reason of the leave
            main_window = browser.window_handles[0]
            leave_link = row.find_element_by_tag_name('a')
            browser.execute_script("arguments[0].scrollIntoView(false);", leave_link)
            leave_link.send_keys(Keys.CONTROL + Keys.RETURN)
            # Edit in new tab
            pop_window = browser.window_handles[1]
            browser.switch_to.window(pop_window)
            element = wait.until(EC.presence_of_element_located((By.TAG_NAME, "tr")))
            for row in browser.find_elements_by_tag_name('tr'):
                # Leave period
                if '差假(或加班)時間' in row.text:
                    lv_dt = row.find_elements_by_tag_name('td')[1].text
                    lv_dt = re.findall(
                        '(\d+-\d+-\d+\(\w+\) \d+:\d+ ~ \d+-\d+-\d+\(\w+\) \d+:\d+)', lv_dt)[0]
                    if searchTomorrow(lv_dt):
                        nextdaylist.append(name)
                elif '事由' in row.text:
                # Leave reason
                    reason = row.find_elements_by_tag_name('td')[1].text
                elif '表單代理人' in row.text:
                # 代理人
                    sub_str = row.find_elements_by_tag_name('td')[1].text
                    sub_unit, sub_title, sub_name = re.findall(
                        '(\w+):\((\w+)\)\(\w+\)(\w+)', sub_str)[0]
                    substitute = sub_unit + sub_name[0] + sub_title                         + sub_name[1:]
                else:
                    continue

            # Close the edit tab
            browser.close()
            browser.switch_to.window(main_window)
            switch_to_iframe(browser)
            leave_str = (titlename_dict[name]+"\t"+leavetype+" "+ lv_dt +
                         '\n\t事由: ' + reason +
                         '\n\t代理人: ' + substitute)
            if titlename_pair[name] not in leavedict:
                leavedict[titlename_pair[name]] = SortedDict()
            if name not in leavedict[titlename_pair[name]]:
                leavedict[titlename_pair[name]][name] = SortedList()
            leavedict[titlename_pair[name]][name].add(leave_str)

print("\r(2/5) 已批核/待批核假單查詢:   完成")


# In[21]:


# 隔日連續假查詢
nextdaylvlist = SortedDict()
sel_day = dt.datetime.today() + dt.timedelta(days = 1)
switch_to_iframe(browser)
print("\r(3/5) 隔日連續假單查詢:   ", end = '')
for name in nextdaylist:
    ## 選取部門
    depart = Select(browser.find_element_by_id("depart"))
    depart.select_by_visible_text(name2office_pair[name])
    # Deselect all members
    try:
        selAllPerson = browser.find_element_by_id('selAll')
        selAllPerson.click()
        selAllPerson.click()
    except:
        pass
    # Select member
    mem_name = NameIdList[name]
    name_xpath = "//input[@value='" + mem_name + "']"
    mem_box = browser.find_element_by_xpath(name_xpath)
    browser.execute_script("arguments[0].scrollIntoView(false);", mem_box)
    mem_box.click()
    ## 選取請假開始日期
    datestrbtn = browser.find_element_by_id("begintime_dc_bt")
    browser.execute_script("arguments[0].scrollIntoView(false);", datestrbtn)
    datestrbtn.click()
    chooseDay(sel_day, browser)
    ##選取請假結束日期
    dateendbtn = browser.find_element_by_id("endtime_dc_bt")
    browser.execute_script("arguments[0].scrollIntoView(false);", dateendbtn)
    dateendbtn.click()
    chooseDay(sel_day, browser)
    for sstate in ["申請已批核完成", "申請未批核完成"]:
        ## 選取表單狀態
        sheetstate = Select(browser.find_element_by_id("condition"))
        sheetstate.select_by_visible_text(sstate)
        ## 點選查詢
        querybtn = browser.find_element_by_name("selecta")
        browser.execute_script("arguments[0].scrollIntoView(false);", querybtn)
        querybtn.click()
        ## 從搜尋結果中擷取長官差假
        table = browser.find_element_by_css_selector("#div2 > table")
        rows = table.find_elements_by_class_name("stripeMe")
        for row in rows:
            leavetype = row.find_elements_by_tag_name("td")[2].text
            # Find the reason of the leave
            main_window = browser.window_handles[0]
            leave_link = row.find_element_by_tag_name('a')
            browser.execute_script("arguments[0].scrollIntoView(false);", leave_link)
            leave_link.send_keys(Keys.CONTROL + Keys.RETURN)
            # Edit in new tab
            pop_window = browser.window_handles[1]
            browser.switch_to.window(pop_window)
            element = wait.until(EC.presence_of_element_located((By.TAG_NAME, "tr")))
            for row in browser.find_elements_by_tag_name('tr'):
                # Leave period
                if '差假(或加班)時間' in row.text:
                    lv_dt = row.find_elements_by_tag_name('td')[1].text
                    lv_dt = re.findall(
                        '(\d+-\d+-\d+\(\w+\) \d+:\d+ ~ \d+-\d+-\d+\(\w+\) \d+:\d+)', lv_dt)[0]
                elif '事由' in row.text:
                # Leave reason
                    reason = row.find_elements_by_tag_name('td')[1].text
                elif '表單代理人' in row.text:
                # 代理人
                    sub_str = row.find_elements_by_tag_name('td')[1].text
                    sub_unit, sub_title, sub_name = re.findall(
                        '(\w+):\((\w+)\)\(\w+\)(\w+)', sub_str)[0]
                    substitute = sub_unit + sub_name[0] + sub_title                         + sub_name[1:]
                else:
                    continue

            # Close the edit tab
            browser.close()
            browser.switch_to.window(main_window)
            switch_to_iframe(browser)
            if not startFromMorning(lv_dt):
                continue
            leave_str = (titlename_dict[name]+"\t明日"+leavetype+"請手動合併 "+ lv_dt +
                         '\n\t事由: ' + reason +
                         '\n\t代理人: ' + substitute)
            if name not in nextdaylvlist:
                nextdaylvlist[name] = SortedList()
            nextdaylvlist[name].add(leave_str)
            
print("\r(3/5) 已批核/待批核假單查詢:   完成")


# In[22]:


# Leave summary
print("\r(3/5) 輸出主管差假清單:   ", end = '')
year = str(dt.date.today().year-1911)
date = dt.date.today().strftime("%m%d")
filename = detail_dir + year + date + "主管差假清單"
filetag = "({}彙整).txt".format(timestamp)
with io.open(filename+filetag, 'w', encoding='utf8') as outfile:
#     outfile.write('\n'.join(leavelist))
    if '局長' in leavedict:
        for name in leavedict['局長']:
            for raw_string in leavedict['局長'][name]:
                outfile.write(raw_string)
                outfile.write('\n')
            if name in nextdaylvlist:
                for raw_string in nextdaylvlist[name]:
                    outfile.write(raw_string)
                    outfile.write('\n')
    if '副局長' in leavedict:
        for name in leavedict['副局長']:
            for raw_string in leavedict['副局長'][name]:
                outfile.write(raw_string)
                outfile.write('\n')
            if name in nextdaylvlist:
                for raw_string in nextdaylvlist[name]:
                    outfile.write(raw_string)
                    outfile.write('\n')
    if '主任秘書' in leavedict:
        for name in leavedict['主任秘書']:
            for raw_string in leavedict['主任秘書'][name]:
                outfile.write(raw_string)
                outfile.write('\n')
            if name in nextdaylvlist:
                for raw_string in nextdaylvlist[name]:
                    outfile.write(raw_string)
                    outfile.write('\n')
    if '簡技' in leavedict:
        for name in leavedict['簡技']:
            for raw_string in leavedict['簡技'][name]:
                outfile.write(raw_string)
                outfile.write('\n')
            if name in nextdaylvlist:
                for raw_string in nextdaylvlist[name]:
                    outfile.write(raw_string)
                    outfile.write('\n')
    if '專委' in leavedict:
        for name in leavedict['專委']:
            for raw_string in leavedict['專委'][name]:
                outfile.write(raw_string)
                outfile.write('\n')
            if name in nextdaylvlist:
                for raw_string in nextdaylvlist[name]:
                    outfile.write(raw_string)
                    outfile.write('\n')
    if '主任' in leavedict:
        for name in leavedict['主任']:
            for raw_string in leavedict['主任'][name]:
                outfile.write(raw_string)
                outfile.write('\n')
            if name in nextdaylvlist:
                for raw_string in nextdaylvlist[name]:
                    outfile.write(raw_string)
                    outfile.write('\n')
    if '科長' in leavedict:
        for name in leavedict['科長']:
            for raw_string in leavedict['科長'][name]:
                outfile.write(raw_string)
                outfile.write('\n')
            if name in nextdaylvlist:
                for raw_string in nextdaylvlist[name]:
                    outfile.write(raw_string)
                    outfile.write('\n')
    if '代理科長' in leavedict:
        for name in leavedict['代理科長']:
            for raw_string in leavedict['代理科長'][name]:
                outfile.write(raw_string)
                outfile.write('\n')
            if name in nextdaylvlist:
                for raw_string in nextdaylvlist[name]:
                    outfile.write(raw_string)
                    outfile.write('\n')

browser.quit()
print("\r(3/5) 輸出主管差假清單:   完成")


# In[23]:


## Update Check
noupdate = False
if len(os.listdir(detail_dir)) > 1:
    past_fn = os.listdir(detail_dir)[-2]
    if year + date in past_fn and        filecmp.cmp(detail_dir + past_fn, filename+filetag):
        noupdate = True
        newtag = filetag.replace('.txt', '_無更改.txt')
        os.rename(filename+filetag, filename+newtag)
        filetag = newtag
if noupdate:
        print("下午無主管差假更新")


# In[24]:


if noupdate:
    print("\r(4/5) 無更新之Word檔")
else:
    print("\r(4/5) 輸出主管差假Word檔:   ", end = '')
    word_tag1 = '({}彙整).docx'.format(timestamp)
    word_tag2 = '_代理人' + word_tag1
    wordfilename1 = file_dir + year + date + '主管差假' + word_tag1
    wordfilename2 = file_dir + year + date + '主管差假' + word_tag2
    # Generate Word file
    doc1 = MyDocument('C:\\Users\\TFD\\主管差假搜尋\\template.docx')
    doc2 = MyDocument('C:\\Users\\TFD\\主管差假搜尋\\template.docx')
    doc1.setTitleDate(dt.date.today())
    doc2.setTitleDate(dt.date.today())
    doc1.writeBigFont('記得更新大隊長差假')
    doc1.writeSmallFont('(私事待辦)')
    if '局長' in leavedict:
        for name in leavedict['局長']:
            for raw_string in leavedict['局長'][name]:
                doc1.addLeave(raw_string, 'all')
                doc2.addLeave(raw_string, 'substitute')
            if name in nextdaylvlist:
                for raw_string in nextdaylvlist[name]:
                    doc1.addLeave(raw_string, 'all')
    if '副局長' in leavedict:
        for name in leavedict['副局長']:
            for raw_string in leavedict['副局長'][name]:
                doc1.addLeave(raw_string, 'all')
                doc2.addLeave(raw_string, 'substitute')
            if name in nextdaylvlist:
                for raw_string in nextdaylvlist[name]:
                    doc1.addLeave(raw_string, 'all')
    if '主任秘書' in leavedict:
        for name in leavedict['主任秘書']:
            for raw_string in leavedict['主任秘書'][name]:
                doc1.addLeave(raw_string, 'all')
                doc2.addLeave(raw_string, 'substitute')
            if name in nextdaylvlist:
                for raw_string in nextdaylvlist[name]:
                    doc1.addLeave(raw_string, 'all')
    if '簡技' in leavedict:
        for name in leavedict['簡技']:
            for raw_string in leavedict['簡技'][name]:
                doc1.addLeave(raw_string, 'all')
                doc2.addLeave(raw_string, 'substitute')
            if name in nextdaylvlist:
                for raw_string in nextdaylvlist[name]:
                    doc1.addLeave(raw_string, 'all')
    if '專委' in leavedict:
        for name in leavedict['專委']:
            for raw_string in leavedict['專委'][name]:
                doc1.addLeave(raw_string, 'all')
                doc2.addLeave(raw_string, 'substitute')
            if name in nextdaylvlist:
                for raw_string in nextdaylvlist[name]:
                    doc1.addLeave(raw_string, 'all')
    if '主任' in leavedict:
        for name in leavedict['主任']:
            for raw_string in leavedict['主任'][name]:
                doc1.addLeave(raw_string, 'all')
            if name in nextdaylvlist:
                for raw_string in nextdaylvlist[name]:
                    doc1.addLeave(raw_string, 'all')
    if '科長' in leavedict:
        for name in leavedict['科長']:
            for raw_string in leavedict['科長'][name]:
                doc1.addLeave(raw_string, 'all')
            if name in nextdaylvlist:
                for raw_string in nextdaylvlist[name]:
                    doc1.addLeave(raw_string, 'all')
    if '代理科長' in leavedict:
        for name in leavedict['代理科長']:
            for raw_string in leavedict['代理科長'][name]:
                doc1.addLeave(raw_string, 'all')
            if name in nextdaylvlist:
                for raw_string in nextdaylvlist[name]:
                    doc1.addLeave(raw_string, 'all')
    doc1.save(wordfilename1)
    doc2.save(wordfilename2)
    print("\r(4/5) 輸出主管差假Word檔:   完成")
    print("主管差假已更新於\""+wordfilename1+"\"")


# In[25]:


if noupdate:
    print("\r(5/5) 無需寄送代理人版本Word檔至主秘室與秘書室")
else:
    print("\r(5/5) 寄送代理人版本Word檔至主秘室與秘書室:   ", end = '')

    import win32com.client as win32

    mail_subject = year + date + ' 簡技以上主管差假'
    mail_subject += '({}彙整)'.format(timestamp)

    outlook = win32.Dispatch('outlook.application')
    mail = outlook.CreateItem(0)
    # mail.To = 'a29536693@tfd.gov.tw'
    mail.To = 'pig_pighead@tfd.gov.tw; jackto@tfd.gov.tw'
    mail.Subject = mail_subject
    mail.Body = '''    您好:
    附件為本日簡技以上主管差假(含代理人與完整請假時間)。

    註: 
    「上午」表示請假期間為當日08-12時，
    「下午」表示請假期間為當日13-17時，
    無特別標註時間表示請假期間為當日08-17時。

    人事室役男  陳定楷
    分機  6726
    '''
    mail.Attachments.Add(wordfilename2)
    mail.Send()

    print("\r(5/5) 寄送代理人版本Word檔至主秘室與秘書室:   完成")


# In[26]:


print("\n\n===主管差假查詢完成===")
if not noupdate:
    print('\n請寄送更新版本之E-mail至局長室/副局長室並電話通知\n')
input("\n按任意鍵結束")

